"""
╔══════════════════════════════════════════════════════════════════╗
║        AUTOMATIZADOR DE JUICIOS DE EVALUACIÓN - SOFÍA PLUS       ║
║                         SENA Colombia                            ║
╚══════════════════════════════════════════════════════════════════╝

REQUISITOS:
    pip install selenium pandas openpyxl webdriver-manager xlrd python-docx

USO:
    python sofia_plus_juicios.py
"""

# ── Stdlib ────────────────────────────────────────────────────────
import os
import sys
import time
import glob
import unicodedata
import re as _re
import io
import threading
import subprocess
import base64
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional, List, Tuple

# ── Habilitar ANSI en Windows ─────────────────────────────────────
if sys.platform == "win32":
    import ctypes
    ctypes.windll.kernel32.SetConsoleMode(
        ctypes.windll.kernel32.GetStdHandle(-11), 7
    )

# ── Dependencias opcionales ───────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

# ── Pandas ────────────────────────────────────────────────────────
import pandas as pd

# ── Selenium ──────────────────────────────────────────────────────
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from webdriver_manager.chrome import ChromeDriverManager

try:
    from webdriver_manager.microsoft import EdgeChromiumManager
except ImportError:
    EdgeChromiumManager = None

# ── Tkinter ───────────────────────────────────────────────────────


# ══════════════════════════════════════════════════════════════════
#  CONFIGURACIÓN GLOBAL
# ══════════════════════════════════════════════════════════════════

URL_LOGIN       = "http://senasofiaplus.edu.co/sofia-public/"
_CARPETA_DEFAULT  = os.environ.get("SOFIA_CARPETA", str(Path.home() / "Downloads" / "SofiaJuicios"))
_CONFIG_FILE      = str(Path.home() / ".sofia_plus_config.txt")
_CHROMEDRIVER_CACHE = str(Path.home() / ".sofia_plus_chromedriver.txt")
_EDGEDRIVER_CACHE   = str(Path.home() / ".sofia_plus_edgedriver.txt")
_LOGIN_CACHE_FILE   = str(Path.home() / ".sofia_plus_login_selector.txt")

ESPERA_MAX    = 12   # segundos
PAUSA_CORTA   = 0.4
MAX_PAGINAS   = 200


# ── Paleta de colores y fuentes (UI) ─────────────────────────────
@dataclass(frozen=True)
class Theme:
    SENA_GREEN  : str = "#39A935"
    SENA_DGREEN : str = "#007832"
    SENA_LGRAY  : str = "#F2F2F2"
    WHITE       : str = "#FFFFFF"
    BG          : str = "#F2F2F2"
    CARD        : str = "#FFFFFF"
    BTN_HOV     : str = "#005A20"
    DANGER      : str = "#C00000"
    LOG_BG      : str = "#1A2B1A"
    LOG_FG      : str = "#E0F0E0"
    TAB_BG      : str = "#D6ECD6"
    F_TITLE     : tuple = ("Arial", 15, "bold")
    F_LABEL     : tuple = ("Arial", 9)
    F_BTN       : tuple = ("Arial", 10, "bold")
    F_LOG       : tuple = ("Consolas", 9)
    F_HINT      : tuple = ("Arial", 8, "italic")

T = Theme()

# ── Consola ANSI ──────────────────────────────────────────────────
VERDE   = "\033[92m"
ROJO    = "\033[91m"
AMARILLO= "\033[93m"
AZUL    = "\033[94m"
RESET   = "\033[0m"
NEGRITA = "\033[1m"

# ── Callback de log (usado por el servidor web para SSE) ─────────
import contextvars as _cv
_log_ctx: _cv.ContextVar = _cv.ContextVar("sofia_log_cb", default=None)

def _emit(level: str, msg: str) -> None:
    cb = _log_ctx.get()
    if cb:
        cb(level, msg)

def ok(msg):
    _emit('ok', msg)
    print(f"  {VERDE}✔{RESET}  {msg}")
def err(msg):
    _emit('err', msg)
    print(f"  {ROJO}✘{RESET}  {msg}")
def info(msg):
    _emit('info', msg)
    print(f"  {AZUL}→{RESET}  {msg}")
def titulo(msg):
    _emit('titulo', msg)
    print(f"\n{NEGRITA}{AMARILLO}{'═'*60}{RESET}")
    print(f"{NEGRITA}{AMARILLO}  {msg}{RESET}")
    print(f"{NEGRITA}{AMARILLO}{'═'*60}{RESET}\n")


# ══════════════════════════════════════════════════════════════════
#  PERSISTENCIA DE CONFIGURACIÓN
# ══════════════════════════════════════════════════════════════════

def _leer_archivo(path: str, default: str = "") -> str:
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return f.read().strip()
    except OSError:
        pass
    return default


def _escribir_archivo(path: str, contenido: str) -> None:
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(contenido)
    except OSError:
        pass


def cargar_carpeta() -> str:
    env = os.environ.get("SOFIA_CARPETA", "").strip()
    if env:
        os.makedirs(env, exist_ok=True)
        return env
    ruta = _leer_archivo(_CONFIG_FILE)
    if ruta and os.path.isdir(ruta):
        return ruta
    return _CARPETA_DEFAULT


def guardar_carpeta(ruta: str) -> None:
    _escribir_archivo(_CONFIG_FILE, ruta)


# ── Estado de la carpeta de descarga (mutable, compartido) ────────
class _Config:
    carpeta_descarga: str = cargar_carpeta()

cfg = _Config()


# ══════════════════════════════════════════════════════════════════
#  DRIVER / NAVEGADOR
# ══════════════════════════════════════════════════════════════════

def _chromium_sistema() -> str:
    """Detecta Chromium instalado en el sistema (Linux/Oracle/Docker/snap)."""
    import shutil
    for binario in ("chromium-browser", "chromium", "chromium-browser-stable"):
        ruta = shutil.which(binario)
        if ruta:
            return ruta
    for ruta in ("/usr/bin/chromium-browser", "/usr/bin/chromium", "/snap/bin/chromium"):
        if os.path.exists(ruta):
            return ruta
    return ""


def _chromedriver_sistema() -> str:
    """Detecta chromedriver instalado en el sistema (Linux/Oracle/Docker/snap)."""
    import shutil
    for binario in ("chromedriver", "chromium.chromedriver"):
        ruta = shutil.which(binario)
        if ruta:
            return ruta
    for ruta in ("/usr/bin/chromedriver", "/snap/bin/chromium.chromedriver"):
        if os.path.exists(ruta):
            return ruta
    return ""


def _chrome_disponible() -> bool:
    import shutil
    if shutil.which("google-chrome") or shutil.which("google-chrome-stable"):
        return True
    if sys.platform == "win32":
        rutas = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        ]
        return any(os.path.exists(r) for r in rutas)
    if sys.platform == "darwin":
        return os.path.exists(
            "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
        )
    return False


def _edge_disponible() -> bool:
    import shutil
    if shutil.which("msedge") or shutil.which("microsoft-edge"):
        return True
    if sys.platform == "win32":
        rutas = [
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        ]
        return any(os.path.exists(r) for r in rutas)
    if sys.platform == "darwin":
        return os.path.exists(
            "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge"
        )
    return False


def _obtener_driver_path(cache_file: str, manager_fn) -> str:
    cached = _leer_archivo(cache_file)
    if cached and os.path.exists(cached):
        return cached
    ruta = manager_fn().install()
    _escribir_archivo(cache_file, ruta)
    return ruta


def _aplicar_opciones_comunes(opts) -> None:
    # NOTA: en Chrome headless --headless=new, los prefs de download son ignorados.
    # La ruta real se fija después vía Browser.setDownloadBehavior (ver configurar_driver).
    # Mantenemos los prefs para compatibilidad con Chrome sin headless / versiones antiguas.
    prefs = {
        "download.default_directory": cfg.carpeta_descarga,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "download.open_pdf_in_system_reader": False,
        "safebrowsing.enabled": False,
        "safebrowsing.disable_download_protection": True,
        "profile.default_content_setting_values.automatic_downloads": 1,
        "profile.default_content_settings.popups": 0,
    }
    opts.add_experimental_option("prefs", prefs)
    for arg in [
        "--headless=new", "--window-size=1920,1080", "--no-sandbox", "--no-zygote",
        "--disable-dev-shm-usage", "--disable-features=TranslateUI,BlinkGenPropertyTrees",
        "--disable-notifications", "--disable-blink-features=AutomationControlled",
        "--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
        "--disable-extensions", "--disable-infobars", "--disable-gpu",
        "--no-first-run", "--no-default-browser-check", "--disable-default-apps",
        "--disable-background-networking", "--disable-sync",
        "--disable-translate", "--safebrowsing-disable-download-protection",
    ]:
        opts.add_argument(arg)
    opts.add_experimental_option(
        "excludeSwitches", ["enable-automation", "enable-logging"]
    )
    opts.add_experimental_option("useAutomationExtension", False)


def configurar_driver():
    os.makedirs(cfg.carpeta_descarga, exist_ok=True)

    chromium_bin    = _chromium_sistema()
    chromedriver_bin = _chromedriver_sistema()

    if chromium_bin and chromedriver_bin:
        info(f"Usando Chromium del sistema: {chromium_bin}")
        opts = webdriver.ChromeOptions()
        opts.binary_location = chromium_bin
        _aplicar_opciones_comunes(opts)
        driver = webdriver.Chrome(
            service=ChromeService(chromedriver_bin),
            options=opts,
        )
    elif _chrome_disponible():
        info("Usando Google Chrome")
        opts = webdriver.ChromeOptions()
        _aplicar_opciones_comunes(opts)
        driver = webdriver.Chrome(
            service=ChromeService(
                _obtener_driver_path(_CHROMEDRIVER_CACHE, ChromeDriverManager)
            ),
            options=opts,
        )
    elif _edge_disponible():
        info("Chrome no encontrado — usando Microsoft Edge")
        if EdgeChromiumManager is None:
            raise RuntimeError(
                "NAVEGADOR_NO_ENCONTRADO: webdriver-manager no incluye soporte para Edge."
            )
        opts = webdriver.EdgeOptions()
        _aplicar_opciones_comunes(opts)
        driver = webdriver.Edge(
            service=EdgeService(
                _obtener_driver_path(_EDGEDRIVER_CACHE, EdgeChromiumManager)
            ),
            options=opts,
        )
    else:
        raise RuntimeError(
            "NAVEGADOR_NO_ENCONTRADO: No se encontró Google Chrome ni Microsoft Edge."
        )

    # Crear carpeta temporal exclusiva para esta sesión de descarga.
    # Esto garantiza que sabemos exactamente dónde Chrome dejará el archivo,
    # sin importar qué archivos viejos haya en cfg.carpeta_descarga.
    import uuid as _uuid
    carpeta_tmp = os.path.join(cfg.carpeta_descarga, f"_tmp_{_uuid.uuid4().hex[:8]}")
    os.makedirs(carpeta_tmp, exist_ok=True)

    # Browser.setDownloadBehavior es el método correcto para Chrome headless --headless=new.
    # Page.setDownloadBehavior está deprecado y es ignorado en headless moderno.
    try:
        driver.execute_cdp_cmd(
            "Browser.setDownloadBehavior",
            {
                "behavior":      "allow",
                "downloadPath":  carpeta_tmp,
                "eventsEnabled": True,
            },
        )
    except Exception:
        # Fallback para versiones antiguas de Chrome/Edge
        driver.execute_cdp_cmd(
            "Page.setDownloadBehavior",
            {"behavior": "allow", "downloadPath": carpeta_tmp},
        )

    # Guardar la carpeta temporal en el driver para que descargar_excel la use
    driver._sofia_tmp_dir = carpeta_tmp
    info(f"Carpeta de descarga temporal: {carpeta_tmp}")
    return driver


# ══════════════════════════════════════════════════════════════════
#  HELPERS DE SELENIUM
# ══════════════════════════════════════════════════════════════════

def _js_clic(driver, elemento) -> None:
    """Clic seguro: intenta clic nativo, con fallback a JS."""
    try:
        elemento.click()
    except Exception:
        driver.execute_script(
            "arguments[0].scrollIntoView({block:'center'}); arguments[0].click();",
            elemento,
        )


def _texto_pagina(driver) -> str:
    """Extrae texto del body del contexto actual."""
    try:
        return driver.execute_script(
            "return document.body ? document.body.innerText.toLowerCase() : '';"
        )
    except Exception:
        return ""


def _texto_completo_pagina_y_frames(driver) -> str:
    """
    Extrae todo el texto visible de la página + iframes en una sola pasada JS.
    Evita switch_to.frame repetidos que son costosos.
    """
    script = """
        var texto = document.body ? document.body.innerText.toLowerCase() : '';
        var iframes = document.querySelectorAll('iframe');
        for (var i = 0; i < iframes.length; i++) {
            try {
                var doc = iframes[i].contentDocument || iframes[i].contentWindow.document;
                if (doc && doc.body) texto += ' ' + doc.body.innerText.toLowerCase();
            } catch(e) {}
        }
        return texto;
    """
    try:
        return driver.execute_script(script) or ""
    except Exception:
        return ""


def _iterar_frames(driver):
    """Generador que yield-ea cada iframe del nivel actual (cambia contexto)."""
    try:
        iframes = driver.find_elements(By.TAG_NAME, "iframe")
    except Exception:
        return
    for i, ifr in enumerate(iframes):
        try:
            driver.switch_to.frame(ifr)
            yield i, ifr
        except Exception:
            pass
        finally:
            try:
                driver.switch_to.default_content()
            except Exception:
                pass


def _pagina_contiene(driver, palabras: List[str]) -> bool:
    """
    Busca palabras en la página y sus iframes con una sola llamada JS.
    Sin switch_to — mucho más rápido que iterar frame a frame.
    """
    try:
        driver.switch_to.default_content()
    except Exception:
        pass
    texto = _texto_completo_pagina_y_frames(driver)
    return any(p in texto for p in palabras)


def _esperar_cambio_dom(driver, firma_antes: int, max_seg: float = 0.8) -> None:
    t0 = time.time()
    while time.time() - t0 < max_seg:
        try:
            firma = driver.execute_script(
                "return document.body ? document.body.innerText.length : 0;"
            )
            if firma != firma_antes:
                return
        except Exception:
            return
        time.sleep(0.05)


def _firma_dom(driver) -> int:
    try:
        return driver.execute_script(
            "return document.body ? document.body.innerText.length : 0;"
        )
    except Exception:
        return 0


# ══════════════════════════════════════════════════════════════════
#  LOGIN
# ══════════════════════════════════════════════════════════════════

# ── XPath único que cubre todos los selectores de usuario en una sola consulta DOM ──
_XPATH_USUARIO = (
    "//input[@name='j_username' or @name='username' or @name='usuario'"
    "       or @id='j_username' or @id='username' or @id='usuario'"
    "       or @type='text'"
    "       or contains(@placeholder,'suario') or contains(@placeholder,'ocumento')]"
    "[not(@type='hidden')][1]"
)
_XPATH_PASSWORD = (
    "//input[@type='password'"
    "       or @name='j_password' or @name='password' or @name='contrasena'"
    "       or @id='j_password' or @id='password'][1]"
)
_XPATH_BOTON = (
    "//input[@type='submit' or @value='Ingresar']"
    " | //button[@type='submit'"
    "            or contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ',"
    "                        'abcdefghijklmnopqrstuvwxyz'),'ingresar')]"
    " | //a[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ',"
    "                'abcdefghijklmnopqrstuvwxyz'),'ingresar')]"
)

# XPath único que detecta el panel post-login O un mensaje de error — una sola llamada
_XPATH_LOGIN_RESULTADO = (
    # Panel OK: elementos típicos del dashboard
    "//*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "             'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'instructor')]"
    " | //*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "                'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'seleccionar rol')]"
    " | //*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "                'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'cerrar sesi')]"
    " | //*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "                'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'bienvenido')]"
    # Error: credenciales incorrectas
    " | //*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "                'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'credenciales')]"
    " | //*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "                'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'contrase')]"
    " | //*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "                'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'acceso denegado')]"
    " | //*[contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
    "                'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'datos incorrectos')]"
)

_PALABRAS_ERROR_LOGIN = {
    "credenciales", "contrase", "acceso denegado", "datos incorrectos",
    "usuario incorrecto", "invalid", "autenticacion", "autenticación",
    "inicio de sesion incorrecto",
}
_PALABRAS_OK_LOGIN = {
    "instructor", "aprendiz", "ejecuci", "formacion", "coordinador",
    "bienvenido", "cerrar sesi", "salir", "perfil", "mis cursos",
    "seleccionar rol",
}


def _encontrar_campo_rapido(driver, xpath: str):
    """
    Busca un campo con un único XPath combinado.
    Sin WebDriverWait: usa find_elements directamente (no bloquea).
    """
    try:
        elems = driver.find_elements(By.XPATH, xpath)
        return elems[0] if elems else None
    except Exception:
        return None


def _rellenar_formulario_login(driver, usuario: str, contrasena: str) -> bool:
    """
    Rellena usuario/contraseña en el contexto actual con un único XPath por campo.
    No usa WebDriverWait — si los campos no están presentes, retorna False al instante.
    """
    campo_usr = _encontrar_campo_rapido(driver, _XPATH_USUARIO)
    if campo_usr is None:
        return False

    campo_pwd = _encontrar_campo_rapido(driver, _XPATH_PASSWORD)
    if campo_pwd is None:
        return False

    # Rellenar via JS para mayor velocidad (evita send_keys carácter a carácter)
    driver.execute_script(
        "arguments[0].value = arguments[1]; "
        "arguments[0].dispatchEvent(new Event('input', {bubbles:true})); "
        "arguments[0].dispatchEvent(new Event('change', {bubbles:true}));",
        campo_usr, usuario,
    )
    driver.execute_script(
        "arguments[0].value = arguments[1]; "
        "arguments[0].dispatchEvent(new Event('input', {bubbles:true})); "
        "arguments[0].dispatchEvent(new Event('change', {bubbles:true}));",
        campo_pwd, contrasena,
    )

    boton = _encontrar_campo_rapido(driver, _XPATH_BOTON)
    if boton:
        try:
            boton.click()
        except Exception:
            driver.execute_script("arguments[0].click();", boton)
    else:
        campo_pwd.send_keys(Keys.RETURN)

    return True



def iniciar_sesion(driver, usuario: str, contrasena: str) -> None:
    titulo("PASO 1: Inicio de sesión en Sofía Plus")
    info("Cargando portal ...")
    driver.get(URL_LOGIN)

    # Esperar solo hasta que aparezca cualquier input (señal mínima de carga)
    try:
        WebDriverWait(driver, ESPERA_MAX, poll_frequency=0.1).until(
            lambda d: d.find_elements(By.XPATH, "//input[@type='text' or @type='password']")
            or d.find_elements(By.TAG_NAME, "iframe")
        )
    except TimeoutException:
        pass

    # ── Estrategia 1: usar el caché del iframe que funcionó la última vez ────
    iframe_cacheado = _leer_archivo(_LOGIN_CACHE_FILE)  # guarda índice como str
    exito = False

    if iframe_cacheado.lstrip("-").isdigit():
        idx = int(iframe_cacheado)
        try:
            if idx == -1:
                # La última vez funcionó sin iframe (página principal)
                exito = _rellenar_formulario_login(driver, usuario, contrasena)
            else:
                iframes = driver.find_elements(By.TAG_NAME, "iframe")
                if idx < len(iframes):
                    driver.switch_to.frame(iframes[idx])
                    exito = _rellenar_formulario_login(driver, usuario, contrasena)
                    driver.switch_to.default_content()
        except Exception:
            driver.switch_to.default_content()
            exito = False

    # ── Estrategia 2: probar página principal primero (más común) ────────────
    if not exito:
        driver.switch_to.default_content()
        exito = _rellenar_formulario_login(driver, usuario, contrasena)
        if exito:
            _escribir_archivo(_LOGIN_CACHE_FILE, "-1")

    # ── Estrategia 3: barrer iframes ─────────────────────────────────────────
    if not exito:
        try:
            iframes = driver.find_elements(By.TAG_NAME, "iframe")
        except Exception:
            iframes = []
        for i, ifr in enumerate(iframes):
            try:
                driver.switch_to.frame(ifr)
                exito = _rellenar_formulario_login(driver, usuario, contrasena)
                if exito:
                    _escribir_archivo(_LOGIN_CACHE_FILE, str(i))
                    break
            except Exception:
                pass
            finally:
                try:
                    driver.switch_to.default_content()
                except Exception:
                    pass

    driver.switch_to.default_content()

    if not exito:
        raise RuntimeError(
            "FORMULARIO_NO_ENCONTRADO: No se detectó el formulario de login. "
            "Verifica que el portal Sofía Plus esté accesible."
        )

    # ── Verificar resultado con una sola llamada JS (sin switch_to por frame) ─
    class _CredencialesInvalidas(Exception):
        pass

    def _verificar_estado(d):
        texto = _texto_completo_pagina_y_frames(d)
        if any(p in texto for p in _PALABRAS_ERROR_LOGIN):
            raise _CredencialesInvalidas()
        return any(p in texto for p in _PALABRAS_OK_LOGIN)

    try:
        WebDriverWait(driver, 25, poll_frequency=0.25).until(_verificar_estado)
        ok("Sesión iniciada correctamente")
    except _CredencialesInvalidas:
        raise RuntimeError(
            "CREDENCIALES_INVALIDAS: El usuario o la contraseña son incorrectos en Sofía Plus."
        )
    except TimeoutException:
        ok("Sesión iniciada — cargando panel principal ...")


# ══════════════════════════════════════════════════════════════════
#  NAVEGACIÓN POR MENÚ (clic con fallback de iframes)
# ══════════════════════════════════════════════════════════════════

# Caché: descripcion → índice de iframe donde se encontró el elemento
_iframe_cache: dict = {}


def _xpath_texto(texto: str) -> str:
    t = texto.lower()
    return (
        f"//*[contains(translate(normalize-space(text()),"
        f"'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
        f"'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'{t}')]"
        f" | //*[contains(translate(@title,"
        f"'ABCDEFGHIJKLMNOPQRSTUVWXYZÁÉÍÓÚÑ',"
        f"'abcdefghijklmnopqrstuvwxyzáéíóúñ'),'{t}')]"
    )


def _buscar_elemento_clickeable(driver, textos: List[str], timeout: float = 1.5):
    """Busca en el contexto actual un elemento clickeable que coincida con algún texto."""
    xpath = " | ".join(_xpath_texto(t) for t in textos)
    # Intento instantáneo primero — si ya está en DOM, no hace falta WebDriverWait
    try:
        elems = driver.find_elements(By.XPATH, xpath)
        if elems and elems[0].is_displayed():
            return elems[0]
    except Exception:
        pass
    try:
        return WebDriverWait(driver, timeout).until(
            EC.element_to_be_clickable((By.XPATH, xpath))
        )
    except Exception:
        return None


def _clic_en_menu(driver, textos: List[str], descripcion: str, timeout: float = 15) -> None:
    """
    Busca y hace clic en un elemento del menú buscando en la página y en todos sus iframes.
    Lanza RuntimeError si no lo encuentra.
    """
    info(f"Buscando: {descripcion} ...")

    def _intentar_en_contexto(t: float):
        elem = _buscar_elemento_clickeable(driver, textos, timeout=t)
        if elem:
            firma = _firma_dom(driver)
            _js_clic(driver, elem)
            _esperar_cambio_dom(driver, firma)
            return True
        return False

    # 1. Intentar desde el caché
    cached_idx = _iframe_cache.get(descripcion)
    if cached_idx is not None:
        try:
            driver.switch_to.default_content()
            iframes = driver.find_elements(By.TAG_NAME, "iframe")
            if cached_idx < len(iframes):
                driver.switch_to.frame(iframes[cached_idx])
                if _intentar_en_contexto(timeout):
                    ok(f"'{descripcion}' seleccionado (iframe cacheado #{cached_idx})")
                    driver.switch_to.default_content()
                    return
        except Exception:
            pass
        finally:
            try:
                driver.switch_to.default_content()
            except Exception:
                pass

    # 2. Página principal
    driver.switch_to.default_content()
    if _intentar_en_contexto(0.3):
        ok(f"'{descripcion}' seleccionado (página principal)")
        return

    # 3. Buscar en iframes y sub-iframes
    try:
        iframes = driver.find_elements(By.TAG_NAME, "iframe")
    except Exception:
        iframes = []

    for i, ifr in enumerate(iframes):
        try:
            driver.switch_to.frame(ifr)
            if _intentar_en_contexto(timeout):
                ok(f"'{descripcion}' seleccionado en iframe #{i}")
                _iframe_cache[descripcion] = i
                driver.switch_to.default_content()
                return
            # Sub-iframes
            try:
                sub_iframes = driver.find_elements(By.TAG_NAME, "iframe")
            except Exception:
                sub_iframes = []
            for sub in sub_iframes:
                try:
                    driver.switch_to.frame(sub)
                    if _intentar_en_contexto(3):
                        ok(f"'{descripcion}' seleccionado en sub-iframe")
                        _iframe_cache[descripcion] = i
                        driver.switch_to.default_content()
                        return
                except Exception:
                    pass
                finally:
                    try:
                        driver.switch_to.parent_frame()
                    except Exception:
                        pass
        except Exception:
            pass
        finally:
            try:
                driver.switch_to.default_content()
            except Exception:
                pass

    raise RuntimeError(
        f"ELEMENTO_NO_ENCONTRADO: No se pudo encontrar '{descripcion}' en la página. "
        "Verifica que la navegación en Sofía Plus haya cargado correctamente."
    )


# ══════════════════════════════════════════════════════════════════
#  PASOS DE NAVEGACIÓN
# ══════════════════════════════════════════════════════════════════

def seleccionar_rol_instructor(driver) -> None:
    titulo("PASO 2: Selección de rol Instructor")
    _clic_en_menu(driver, ["instructor", "Instructor"], "Rol Instructor", timeout=8)


def navegar_reporte_juicios(driver) -> None:
    titulo("PASO 3: Navegación al reporte de juicios")
    _clic_en_menu(
        driver,
        ["ejecución de la formación", "ejecucion de la formacion", "ejecución formación"],
        "Ejecución de la Formación", timeout=8,
    )
    _clic_en_menu(
        driver,
        ["administrar ruta de aprendizaje", "administrar ruta"],
        "Administrar Ruta de Aprendizaje", timeout=8,
    )
    _clic_en_menu(driver, ["reportes", "reporte"], "Reportes", timeout=8)
    _clic_en_menu(
        driver,
        ["reporte de juicios de evaluación", "reporte de juicios", "juicios de evaluación"],
        "Reporte de Juicios de Evaluación", timeout=8,
    )
    ok("Navegación completada — estamos en el Reporte de Juicios de Evaluación")


# ══════════════════════════════════════════════════════════════════
#  BÚSQUEDA DE FICHA Y DESCARGA
# ══════════════════════════════════════════════════════════════════

_JS_BUSCAR_ICONO_FICHA = """
    var fid = arguments[0];
    var links = document.querySelectorAll('a[onclick]');
    for (var i = 0; i < links.length; i++) {
        if ((links[i].getAttribute('onclick') || '').indexOf(fid) !== -1) {
            return links[i].querySelector('img') || links[i];
        }
    }
    var cells = document.querySelectorAll('td, span');
    for (var j = 0; j < cells.length; j++) {
        if (cells[j].textContent.trim() === fid) {
            var row = cells[j].closest('tr');
            if (row) {
                return row.querySelector('img[src*="seleccionar"]') || row.querySelector('a');
            }
        }
    }
    return null;
"""


def _entrar_frame_popup(driver) -> bool:
    """
    Busca y entra al iframe que contiene el popup de fichas.
    Devuelve True si lo encontró.
    """
    driver.switch_to.default_content()
    try:
        iframes = driver.find_elements(By.TAG_NAME, "iframe")
    except Exception:
        return False

    for ifr in iframes:
        try:
            driver.switch_to.frame(ifr)
            try:
                driver.find_element(
                    By.XPATH,
                    "//form[@id='form2'] | //a[contains(@id,'dsListas')]",
                )
                return True
            except Exception:
                pass
            # Intentar sub-iframes
            for sub in driver.find_elements(By.TAG_NAME, "iframe"):
                try:
                    driver.switch_to.frame(sub)
                    try:
                        driver.find_element(
                            By.XPATH,
                            "//form[@id='form2'] | //a[contains(@id,'dsListas')]",
                        )
                        return True
                    except Exception:
                        pass
                    driver.switch_to.parent_frame()
                except Exception:
                    driver.switch_to.parent_frame()
        except Exception:
            pass
        driver.switch_to.default_content()

    return False


def _icono_ficha_en_pagina(driver, numero_ficha: str):
    """Devuelve el elemento de selección de la ficha si está en la página actual."""
    try:
        elem = driver.execute_script(_JS_BUSCAR_ICONO_FICHA, numero_ficha)
        if elem:
            return elem
    except Exception:
        pass
    xpath = (
        f"//a[contains(@onclick,'{numero_ficha}')]//img | "
        f"//a[contains(@onclick,'{numero_ficha}')] | "
        f"//tr[.//*[normalize-space(.)='{numero_ficha}']]//img[contains(@src,'seleccionar')] | "
        f"//tr[.//*[normalize-space(.)='{numero_ficha}']]//a[1]"
    )
    try:
        elems = driver.find_elements(By.XPATH, xpath)
        return elems[0] if elems else None
    except Exception:
        return None


def _boton_siguiente_pagina(driver):
    """Devuelve el botón 'siguiente página' si existe y está visible."""
    for btn_id in ["form2:dsListasnext", "dsListasnext", "nextPage", "btnNext"]:
        elems = driver.find_elements(By.ID, btn_id)
        if elems and elems[0].is_displayed():
            return elems[0]
    for xpath in [
        "//a[contains(@id,'next') or contains(@id,'Next') or contains(@id,'siguiente')]",
        "//img[contains(@src,'next') or contains(@src,'siguiente')]/..",
        "//a[contains(@onclick,'next') or contains(@title,'Siguiente')]",
    ]:
        elems = driver.find_elements(By.XPATH, xpath)
        if elems and elems[0].is_displayed():
            return elems[0]
    return None


def _contenido_tabla_actual(driver) -> str:
    """Extrae el HTML de la tabla para detectar si la página cambió."""
    try:
        return driver.execute_script(
            "var t=document.querySelector('table'); return t ? t.innerHTML : '';"
        )
    except Exception:
        return ""


def _buscar_ficha_paginando(driver, numero_ficha: str) -> bool:
    """
    Recorre todas las páginas del popup de fichas buscando la ficha indicada.
    Devuelve True si la encontró y la seleccionó.
    """
    if not _entrar_frame_popup(driver):
        # Último intento: buscar directamente sin entrar a un frame
        info("Frame del popup no encontrado, buscando en contexto actual ...")

    contenido_anterior: Optional[str] = None

    for pagina in range(1, MAX_PAGINAS + 1):
        if pagina > 1:
            _entrar_frame_popup(driver)

        icono = _icono_ficha_en_pagina(driver, numero_ficha)
        if icono:
            ok(f"Ficha {numero_ficha} encontrada en página {pagina}")
            _js_clic(driver, icono)
            ok(f"Ficha {numero_ficha} seleccionada")
            driver.switch_to.default_content()
            return True

        contenido_actual = _contenido_tabla_actual(driver)
        if contenido_anterior is not None and contenido_actual == contenido_anterior:
            info("El contenido no cambió — última página alcanzada.")
            break
        contenido_anterior = contenido_actual

        btn_next = _boton_siguiente_pagina(driver)
        if btn_next is None:
            info("Sin botón siguiente — última página.")
            break

        html_antes = contenido_actual
        _js_clic(driver, btn_next)
        t0 = time.time()
        while time.time() - t0 < 1.5:
            if _contenido_tabla_actual(driver) != html_antes:
                break
            time.sleep(0.04)

    driver.switch_to.default_content()
    return False


def buscar_ficha_y_aprendiz(driver, numero_ficha: str, mtime_previo_ref: list) -> None:
    """
    Abre el popup de fichas, selecciona la ficha y hace clic en Consultar.

    Args:
        mtime_previo_ref: lista de un elemento [float] donde se almacenará
                          el mtime del XLS justo antes del clic Consultar,
                          para detectar correctamente un nuevo archivo.
    """
    titulo("PASO 4: Selección de ficha")

    # Abrir popup
    info("Abriendo popup de fichas ...")
    try:
        iframes = driver.find_elements(By.TAG_NAME, "iframe")
        driver.switch_to.frame(iframes[0])
        btn_popup = WebDriverWait(driver, ESPERA_MAX).until(
            EC.element_to_be_clickable(
                (By.XPATH, "//img[contains(@src,'consultar_en_pop_up')]")
            )
        )
        _js_clic(driver, btn_popup)
        driver.switch_to.default_content()
        ok("Popup abierto")
    except Exception as ex:
        driver.switch_to.default_content()
        raise RuntimeError(
            f"POPUP_NO_ABIERTO: No se pudo abrir el popup de fichas. Detalle: {ex}"
        )

    # Esperar que el popup cargue
    try:
        WebDriverWait(driver, 5, poll_frequency=0.15).until(
            lambda d: d.execute_script(
                "var iframes=document.querySelectorAll('iframe');"
                "for(var i=0;i<iframes.length;i++){"
                "  try{var d2=iframes[i].contentDocument||iframes[i].contentWindow.document;"
                "  if(d2&&d2.querySelector('table,a[onclick]'))return true;}catch(e){}"
                "}return false;"
            )
        )
    except TimeoutException:
        pass

    info(f"Buscando ficha {numero_ficha} en el popup (con paginación) ...")
    seleccionada = _buscar_ficha_paginando(driver, numero_ficha)

    if not seleccionada:
        raise RuntimeError(
            f"FICHA_NO_ENCONTRADA: No se encontró la ficha {numero_ficha} "
            "en el listado de Sofía Plus. Verifica el número de ficha."
        )

    # La carpeta temporal es exclusiva de esta sesión, siempre estará vacía antes del clic.
    # Usamos mtime 0 — cualquier archivo que aparezca es el que queremos.
    mtime_previo_ref[0] = 0

    # Hacer clic en "Consultar"
    info("Haciendo clic en 'Consultar' ...")
    try:
        def _consultar_listo(d):
            try:
                iframes = d.find_elements(By.TAG_NAME, "iframe")
                if not iframes:
                    return False
                d.switch_to.frame(iframes[0])
                btn = d.find_element(By.ID, "frmForma1:btnConsultar")
                return btn if btn.is_enabled() else False
            except Exception:
                return False
            finally:
                try:
                    d.switch_to.default_content()
                except Exception:
                    pass

        WebDriverWait(driver, 8, poll_frequency=0.2).until(_consultar_listo)
        iframes = driver.find_elements(By.TAG_NAME, "iframe")
        driver.switch_to.frame(iframes[0])
        btn_consultar = driver.find_element(By.ID, "frmForma1:btnConsultar")
        _js_clic(driver, btn_consultar)
        driver.switch_to.default_content()
        ok("Reporte generado")
    except Exception as ex:
        driver.switch_to.default_content()
        raise RuntimeError(
            f"CONSULTAR_FALLIDO: No se pudo hacer clic en 'Consultar'. Detalle: {ex}"
        )

    driver.switch_to.default_content()

    # Espera reactiva: vigila la carpeta temporal del driver
    _tmp = getattr(driver, "_sofia_tmp_dir", cfg.carpeta_descarga)
    t0 = time.time()
    while time.time() - t0 < 10:
        if glob.glob(os.path.join(_tmp, "*.crdownload")) or glob.glob(os.path.join(_tmp, "*.xls*")):
            break
        time.sleep(0.05)


# ══════════════════════════════════════════════════════════════════
#  DESCARGA DEL EXCEL
# ══════════════════════════════════════════════════════════════════

def descargar_excel(
    numero_ficha: str,
    mtime_previo: float,
    t_inicio: float = 0.0,
    carpeta_tmp: str = "",
) -> Optional[str]:
    """
    Espera el XLS en `carpeta_tmp` (carpeta temporal exclusiva creada por configurar_driver),
    lo renombra y lo mueve a cfg.carpeta_descarga/Ficha_XXXXX/.
    Elimina la carpeta temporal al finalizar.
    """
    titulo("PASO 5: Localizando archivo descargado")

    # Marcar inicio si no viene del caller (para medir tiempo total)
    _t0 = t_inicio if t_inicio > 0 else time.time()

    # Si no se pasó carpeta_tmp, usar la carpeta configurada directamente
    directorio = carpeta_tmp if carpeta_tmp and os.path.isdir(carpeta_tmp) else cfg.carpeta_descarga
    info(f"Esperando archivo en: {directorio}")

    # ── Esperar que aparezca el archivo XLS completo ─────────────
    # Un archivo está completo cuando:
    #   1. No hay .crdownload en la carpeta (descarga en progreso)
    #   2. Existe un .xls* que no sea parcial
    #   3. Su tamaño no cambia en dos lecturas consecutivas (Chrome lo liberó)
    archivo_descargado: Optional[str] = None
    t_lim = time.time() + 90  # hasta 90 s para conexiones lentas

    _ultimo_candidato = ""
    _ultimo_size      = -1

    while time.time() < t_lim:
        # Esperar que desaparezca cualquier .crdownload
        if glob.glob(os.path.join(directorio, "*.crdownload")):
            time.sleep(0.1)
            continue

        candidatos = [
            f for f in glob.glob(os.path.join(directorio, "*.xls*"))
            if not f.endswith(".crdownload") and not f.endswith(".tmp")
        ]
        if not candidatos:
            time.sleep(0.1)
            continue

        candidato = max(candidatos, key=os.path.getmtime)
        try:
            size_actual = os.path.getsize(candidato)
        except OSError:
            time.sleep(0.1)
            continue

        # Si el tamaño estabilizó respecto a la lectura anterior → listo
        if candidato == _ultimo_candidato and size_actual == _ultimo_size and size_actual > 0:
            archivo_descargado = candidato
            break

        _ultimo_candidato = candidato
        _ultimo_size      = size_actual
        time.sleep(0.15)  # esperar un poco antes de re-leer el tamaño

    if not archivo_descargado:
        err(f"El archivo no apareció en {directorio} después de 90 s.")
        err("Verifica que el reporte exista para esta ficha en Sofía Plus.")
        _limpiar_tmp(directorio, carpeta_tmp)
        return None

    _t_descarga = time.time() - _t0
    ok(f"Archivo descargado: {os.path.basename(archivo_descargado)} ({_t_descarga:.1f} s)")

    # ── Renombrar y mover a cfg.carpeta_descarga/Ficha_XXXXX/ ────
    ext          = os.path.splitext(archivo_descargado)[1]
    destino_dir  = _carpeta_ficha(numero_ficha)
    nombre_final = f"Reporte de Juicios Evaluativos - {numero_ficha}{ext}"
    ruta_final   = os.path.join(destino_dir, nombre_final)

    # Reintentar el rename hasta 3 s por si Chrome aún tiene el handle abierto
    _movido = False
    _t_rename = time.time() + 3
    while time.time() < _t_rename:
        try:
            if os.path.exists(ruta_final):
                os.remove(ruta_final)
            os.rename(archivo_descargado, ruta_final)
            _movido = True
            break
        except OSError:
            time.sleep(0.2)  # Chrome puede tener el handle abierto brevemente

    if not _movido:
        err(f"No se pudo mover el archivo a {destino_dir}")
        _limpiar_tmp(directorio, carpeta_tmp)
        return None

    _t_total = time.time() - _t0
    ok(f"Guardado en: {os.path.relpath(ruta_final, cfg.carpeta_descarga)}")
    info(f"Tiempo total (descarga + mover): {_t_total:.1f} s")

    _limpiar_tmp(directorio, carpeta_tmp)
    return ruta_final


def _limpiar_tmp(directorio: str, carpeta_tmp: str) -> None:
    """Elimina completamente la carpeta temporal y todo su contenido."""
    if not carpeta_tmp or not os.path.isdir(carpeta_tmp):
        return
    if carpeta_tmp == cfg.carpeta_descarga:
        return  # nunca borrar la carpeta principal
    # Verificar que sea una subcarpeta de cfg.carpeta_descarga (seguridad extra)
    try:
        if not os.path.abspath(carpeta_tmp).startswith(os.path.abspath(cfg.carpeta_descarga)):
            return
    except Exception:
        return
    try:
        import shutil
        shutil.rmtree(carpeta_tmp, ignore_errors=True)
        info(f"Carpeta temporal eliminada: {os.path.basename(carpeta_tmp)}")
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════
#  ANÁLISIS DEL EXCEL
# ══════════════════════════════════════════════════════════════════

def _sin_tildes(texto: str) -> str:
    return unicodedata.normalize("NFD", texto).encode("ascii", "ignore").decode("utf-8")


def _es_pendiente(valor) -> bool:
    v = str(valor).strip().lower()
    return "por evaluar" in v or "no aprobado" in v or "no aprobad" in v


@dataclass
class _ColumnasMeta:
    nombre: str
    apellido: str
    estado: str
    competencia: str
    ra: str
    juicio: str


def _parsear_excel(ruta_excel: str) -> Tuple[pd.DataFrame, _ColumnasMeta]:
    """
    Lee el Excel de Sofía Plus y devuelve el DataFrame y los nombres de columnas.
    Centraliza toda la lógica de lectura/normalización del archivo.
    """
    try:
        df = pd.read_excel(ruta_excel, header=None, engine="xlrd")
    except Exception:
        df = pd.read_excel(ruta_excel, header=None)

    df.columns = df.iloc[12].astype(str).str.strip().tolist()
    df = df.iloc[13:].copy()
    df.reset_index(drop=True, inplace=True)
    df.dropna(how="all", inplace=True)

    cols = df.columns.tolist()
    meta = _ColumnasMeta(
        nombre=cols[2], apellido=cols[3], estado=cols[4],
        competencia=cols[5], ra=cols[6], juicio=cols[7],
    )
    return df, meta


def _filtrar_aprendices(
    df: pd.DataFrame,
    meta: _ColumnasMeta,
    nombre: str,
    apellido: str,
    solo_activos: bool,
) -> pd.DataFrame:
    """Aplica los filtros de nombre, apellido y estado activo."""
    mask = pd.Series([True] * len(df))

    for palabra in [p.upper() for p in nombre.split() if len(p) >= 2]:
        mask &= df[meta.nombre].astype(str).str.upper().str.contains(palabra, na=False)
    for palabra in [p.upper() for p in apellido.split() if len(p) >= 2]:
        mask &= df[meta.apellido].astype(str).str.upper().str.contains(palabra, na=False)

    if not nombre and not apellido and solo_activos:
        ESTADOS_ACTIVOS = ["en formacion", "condicionado"]
        mask_activos = df[meta.estado].astype(str).str.strip().apply(
            lambda v: any(ea in _sin_tildes(v.lower()) for ea in ESTADOS_ACTIVOS)
        )
        mask &= mask_activos
        info(f"Filtro activos aplicado — {mask.sum()} filas encontradas")

    return df[mask].copy()


def _crear_doc_base(criterio_desc: str, total_aprendices: int) -> "DocxDocument":
    doc = DocxDocument()
    t = doc.add_heading("Informe de Juicios de Evaluacion - Sofia Plus", level=1)
    t.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    for lbl, val in [
        ("Criterio de búsqueda:", criterio_desc),
        ("Generado:", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Total aprendices:", str(total_aprendices)),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{lbl} ").bold = True
        p.add_run(val)
    return doc


def _tabla_resumen_doc(doc, total_ap: int, aprobados_ap: int,
                       pend_por_eval: int, pend_no_aprob: int) -> None:
    tabla = doc.add_table(rows=4, cols=2)
    tabla.style = "Table Grid"
    filas = [
        ("Total RA",     str(total_ap)),
        ("Aprobados",    str(aprobados_ap)),
        ("Por Evaluar",  str(pend_por_eval)),
        ("No Aprobados", str(pend_no_aprob)),
    ]
    for i, (lbl, val) in enumerate(filas):
        tabla.rows[i].cells[0].paragraphs[0].add_run(lbl).bold = True
        run = tabla.rows[i].cells[1].paragraphs[0].add_run(val)
        if lbl == "No Aprobados" and int(val) > 0:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif lbl == "Por Evaluar" and int(val) > 0:
            run.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)


def _secciones_pendientes_doc(doc, df_pend: pd.DataFrame, meta: _ColumnasMeta) -> None:
    doc.add_paragraph()
    doc.add_heading("Resultados Pendientes", level=3)
    for comp, grupo in df_pend.groupby(meta.competencia, sort=False):
        comp_str = str(comp).strip()
        if comp_str.lower() in ("nan", ""):
            comp_str = "Sin competencia asignada"
        h = doc.add_heading(comp_str, level=4)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
        for _, fila in grupo.iterrows():
            ra  = str(fila[meta.ra]).strip()
            jui = str(fila[meta.juicio]).strip()
            p = doc.add_paragraph(style="List Number")
            p.add_run(ra)
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Pt(24)
            p2.add_run("Estado: ").bold = True
            run_j = p2.add_run(jui)
            run_j.font.color.rgb = (
                RGBColor(0xC0, 0x00, 0x00)
                if "no aprobado" in jui.lower()
                else RGBColor(0xFF, 0x80, 0x00)
            )


def analizar_juicios(
    ruta_excel: str,
    nombre: str,
    apellido: str,
    solo_activos: bool = False,
) -> None:
    """Genera un documento DOCX consolidado con todos los aprendices filtrados."""
    _t_inf0 = time.time()
    titulo("PASO 6: Análisis de resultados de aprendizaje")
    info(f"Procesando: {os.path.basename(ruta_excel)}")

    nombre   = (nombre   or "").strip()
    apellido = (apellido or "").strip()

    df, meta = _parsear_excel(ruta_excel)
    df_filtrado = _filtrar_aprendices(df, meta, nombre, apellido, solo_activos)

    if df_filtrado.empty:
        criterio = (nombre + " " + apellido).strip() or "todos"
        err(f"No se encontraron aprendices con criterio: '{criterio}'")
        _mostrar_muestra(df, meta)
        return

    aprendices_unicos = df_filtrado[[meta.nombre, meta.apellido]].drop_duplicates()
    total = len(aprendices_unicos)
    criterio_desc = _describir_criterio(nombre, apellido, aprendices_unicos, meta)
    ok(f"{total} aprendice(s) encontrado(s): {criterio_desc}")

    sufijo = _sufijo_nombre_archivo(nombre, apellido)
    ruta_resumen = os.path.join(
        os.path.dirname(ruta_excel),
        f"RESUMEN_{sufijo}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
    )

    if not DOCX_DISPONIBLE:
        err("python-docx no instalado. Instala con: python -m pip install python-docx")
        return

    doc = _crear_doc_base(criterio_desc, total)
    sep = "=" * 62

    for idx, (_, fila_ap) in enumerate(aprendices_unicos.iterrows(), 1):
        nom_ap = str(fila_ap[meta.nombre]).strip()
        ape_ap = str(fila_ap[meta.apellido]).strip()
        nombre_completo = f"{nom_ap} {ape_ap}"

        mask_ap = (
            (df_filtrado[meta.nombre].astype(str).str.strip() == nom_ap)
            & (df_filtrado[meta.apellido].astype(str).str.strip() == ape_ap)
        )
        df_ap    = df_filtrado[mask_ap]
        df_pend  = df_ap[df_ap[meta.juicio].apply(_es_pendiente)]
        df_aprob = df_ap[~df_ap[meta.juicio].apply(_es_pendiente)]

        total_ap     = len(df_ap)
        aprobados_ap = len(df_aprob)
        pendientes_ap = len(df_pend)

        # Consola
        print(f"\n{NEGRITA}{sep}{RESET}")
        print(f"{NEGRITA}  {idx}/{total} — {nombre_completo.upper()}{RESET}")
        print(f"{NEGRITA}{sep}{RESET}")
        print(f"  {'Total RA:'.ljust(34)}{NEGRITA}{total_ap}{RESET}")
        print(f"  {VERDE}{'Aprobados:'.ljust(34)}{RESET}{NEGRITA}{aprobados_ap}{RESET}")
        print(f"  {ROJO}{'Pendientes:'.ljust(34)}{RESET}{NEGRITA}{pendientes_ap}{RESET}")
        print()

        if pendientes_ap > 0:
            for comp, grupo in df_pend.groupby(meta.competencia, sort=False):
                comp_str = str(comp).strip()
                if comp_str.lower() in ("nan", ""):
                    comp_str = "Sin competencia asignada"
                print(f"  {AMARILLO}{NEGRITA}Competencia: {comp_str}{RESET}")
                for i2, (_, f2) in enumerate(grupo.iterrows(), 1):
                    ra  = str(f2[meta.ra]).strip()
                    jui = str(f2[meta.juicio]).strip()
                    color = ROJO if "no aprobado" in jui.lower() else AMARILLO
                    print(f"    {i2}. {ra}")
                    print(f"       Estado: {color}{jui}{RESET}")
                print()
        else:
            print(f"  {VERDE}{NEGRITA}¡Todos los resultados aprobados!{RESET}\n")

        # DOCX
        doc.add_paragraph()
        h2 = doc.add_heading(nombre_completo.upper(), level=2)
        h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        pend_no_aprob = len(
            df_pend[
                df_pend[meta.juicio].astype(str).str.lower().str.contains("no aprobado", na=False)
            ]
        )
        pend_por_eval = pendientes_ap - pend_no_aprob
        _tabla_resumen_doc(doc, total_ap, aprobados_ap, pend_por_eval, pend_no_aprob)

        if pendientes_ap > 0:
            _secciones_pendientes_doc(doc, df_pend, meta)
        else:
            p = doc.add_paragraph()
            r = p.add_run("Todos los resultados aprobados.")
            r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
            r.bold = True

    doc.save(ruta_resumen)
    _t_inf = time.time() - _t_inf0
    ok(f"Resumen guardado: {os.path.basename(ruta_resumen)}")
    print(f"  {AMARILLO}{NEGRITA}⏱  Generación de informe(s): {_t_inf:.1f} s{RESET}")
    return ruta_resumen


def analizar_juicios_independientes(
    ruta_excel: str,
    nombre: str,
    apellido: str,
    numero_ficha: str,
    solo_activos: bool = False,
) -> None:
    """Genera un DOCX individual por aprendiz en una subcarpeta."""
    _t_inf0 = time.time()
    titulo("Generando reportes independientes por aprendiz")
    info(f"Procesando: {os.path.basename(ruta_excel)}")

    if not DOCX_DISPONIBLE:
        err("python-docx no disponible. Instala con: python -m pip install python-docx")
        return

    nombre   = (nombre   or "").strip()
    apellido = (apellido or "").strip()

    df, meta = _parsear_excel(ruta_excel)
    df_filtrado = _filtrar_aprendices(df, meta, nombre, apellido, solo_activos)

    if df_filtrado.empty:
        criterio = (nombre + " " + apellido).strip() or "todos"
        err(f"No se encontraron aprendices con criterio: {criterio}")
        return

    aprendices = df_filtrado[[meta.nombre, meta.apellido]].drop_duplicates()
    total = len(aprendices)
    ok(f"{total} aprendice(s) encontrado(s)")

    timestamp      = datetime.now().strftime("%Y%m%d_%H%M")
    carpeta_salida = os.path.join(_carpeta_ficha(numero_ficha), f"Individuales_{timestamp}")
    os.makedirs(carpeta_salida, exist_ok=True)
    ok(f"Carpeta de reportes: Ficha_{numero_ficha}/Individuales_{timestamp}")

    for idx, (_, fila_ap) in enumerate(aprendices.iterrows(), 1):
        nom_ap  = str(fila_ap[meta.nombre]).strip()
        ape_ap  = str(fila_ap[meta.apellido]).strip()
        nombre_completo = f"{nom_ap} {ape_ap}"

        mask_ap = (
            (df_filtrado[meta.nombre].astype(str).str.strip() == nom_ap)
            & (df_filtrado[meta.apellido].astype(str).str.strip() == ape_ap)
        )
        df_ap    = df_filtrado[mask_ap]
        df_pend  = df_ap[df_ap[meta.juicio].apply(_es_pendiente)]
        df_aprob = df_ap[~df_ap[meta.juicio].apply(_es_pendiente)]

        total_ap     = len(df_ap)
        aprobados_ap = len(df_aprob)
        pend_ap      = len(df_pend)

        info(f"[{idx}/{total}] {nombre_completo}")

        doc = DocxDocument()
        t = doc.add_heading("Informe de Juicios de Evaluacion - Sofia Plus", level=1)
        t.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        doc.add_heading("Datos del Aprendiz", level=2)
        for lbl, val in [
            ("Aprendiz:", nombre_completo.upper()),
            ("Ficha:", numero_ficha),
            ("Generado:", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{lbl} ").bold = True
            p.add_run(val)

        doc.add_heading("Resumen", level=2)
        pend_no_aprob = len(
            df_pend[
                df_pend[meta.juicio].astype(str).str.lower().str.contains("no aprobado", na=False)
            ]
        )
        pend_por_eval = pend_ap - pend_no_aprob
        _tabla_resumen_doc(doc, total_ap, aprobados_ap, pend_por_eval, pend_no_aprob)

        if pend_ap > 0:
            doc.add_heading("Resultados Pendientes por Competencia", level=2)
            _secciones_pendientes_doc(doc, df_pend, meta)
        else:
            p = doc.add_paragraph()
            r = p.add_run("Todos los resultados aprobados.")
            r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
            r.bold = True

        nombre_arch = f"{ape_ap}_{nom_ap}".replace(" ", "_") + ".docx"
        doc.save(os.path.join(carpeta_salida, nombre_arch))
        ok(f"Guardado: {nombre_arch}  (Pend: {pend_ap}/{total_ap})")

    _t_inf = time.time() - _t_inf0
    ok(f"Proceso completado — {total} archivo(s) en: {carpeta_salida}")
    return carpeta_salida
    print(f"  {AMARILLO}{NEGRITA}⏱  Generación de informe(s): {_t_inf:.1f} s{RESET}")


# ── Helpers de análisis ───────────────────────────────────────────

def _mostrar_muestra(df: pd.DataFrame, meta: _ColumnasMeta) -> None:
    info("Muestra de nombres en el Excel:")
    for _, fila in df[[meta.nombre, meta.apellido]].drop_duplicates().head(10).iterrows():
        print(f"     {fila[meta.nombre]} | {fila[meta.apellido]}")


def _describir_criterio(
    nombre: str, apellido: str, aprendices_unicos: pd.DataFrame, meta: _ColumnasMeta
) -> str:
    if len(aprendices_unicos) == 1:
        f = aprendices_unicos.iloc[0]
        return f"{f[meta.nombre]} {f[meta.apellido]}"
    if not nombre and not apellido:
        return "Todos los aprendices de la ficha"
    if nombre and not apellido:
        return f"Aprendices con nombre: {nombre}"
    if apellido and not nombre:
        return f"Aprendices con apellido: {apellido}"
    return f"{nombre} {apellido}"


def _sufijo_nombre_archivo(nombre: str, apellido: str) -> str:
    if not nombre and not apellido:
        return "TODOS"
    if nombre and not apellido:
        return f"nombre_{nombre.replace(' ', '_')}"
    if apellido and not nombre:
        return f"apellido_{apellido.replace(' ', '_')}"
    return f"{apellido.replace(' ', '_')}_{nombre.replace(' ', '_')}"


# ══════════════════════════════════════════════════════════════════
#  UTILIDADES DE ARCHIVOS
# ══════════════════════════════════════════════════════════════════

def _carpeta_ficha(numero_ficha: str) -> str:
    ruta = os.path.join(cfg.carpeta_descarga, f"Ficha_{numero_ficha}")
    os.makedirs(ruta, exist_ok=True)
    return ruta


def buscar_excel_ficha(numero_ficha: str) -> Optional[str]:
    for ext in (".xls", ".xlsx"):
        ruta = os.path.join(
            cfg.carpeta_descarga, f"Ficha_{numero_ficha}",
            f"Reporte de Juicios Evaluativos - {numero_ficha}{ext}",
        )
        if os.path.exists(ruta):
            return ruta
    subcarpeta = os.path.join(cfg.carpeta_descarga, f"Ficha_{numero_ficha}")
    if os.path.isdir(subcarpeta):
        for f in glob.glob(os.path.join(subcarpeta, "*.xls*")):
            return f
    for f in glob.glob(os.path.join(cfg.carpeta_descarga, "*.xls*")):
        if numero_ficha in os.path.basename(f):
            return f
    return None


# ══════════════════════════════════════════════════════════════════
#  INTERFAZ GRÁFICA (tkinter)
# ══════════════════════════════════════════════════════════════════

